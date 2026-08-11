from __future__ import annotations

import re
from copy import deepcopy
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE_DIR = Path(__file__).resolve().parent.parent
SPECIFICATION_TEMPLATE = BASE_DIR / "document_templates" / "machine_specification_template.docx"
SECTION_KEYS = ("system", "connectivity", "warranty")
PRICE_TERMS = {"price", "pricing", "cost", "rrp", "msrp", "unit price", "total price"}


def _clean(value: object, limit: int = 1500) -> str:
    return " ".join(str(value or "").split())[:limit]


def _contains_price(label: str) -> bool:
    lowered = label.casefold()
    return any(term in lowered for term in PRICE_TERMS)


def _contains_price_value(value: str) -> bool:
    return bool(
        re.search(r"[£€$]\s*\d", value)
        or re.search(r"\b(?:GBP|USD|EUR)\s*\d", value, flags=re.IGNORECASE)
        or re.search(r"\d\s*(?:GBP|USD|EUR)\b", value, flags=re.IGNORECASE)
        or re.search(r"\b(?:price|cost|RRP|MSRP)\b", value, flags=re.IGNORECASE)
    )


def _normalise_sections(parsed: dict, sources: list[dict]) -> dict[str, list[dict[str, str]]]:
    sections: dict[str, list[dict[str, str]]] = {key: [] for key in SECTION_KEYS}
    raw_sections = parsed.get("sections") if isinstance(parsed, dict) else None
    if not isinstance(raw_sections, dict):
        raise ValueError("The model did not return specification sections.")
    for key in SECTION_KEYS:
        rows = raw_sections.get(key) or []
        if not isinstance(rows, list):
            continue
        for row in rows[:30]:
            if not isinstance(row, dict):
                continue
            label = _clean(row.get("label"), 120)
            value = _clean(row.get("value"))
            if not label or not value or _contains_price(label) or _contains_price_value(value):
                continue
            sections[key].append({"label": label, "value": value, "url": ""})

    seen_urls: set[str] = set()
    for source in sources[:12]:
        url = _clean(source.get("url"), 2000)
        if not url or url in seen_urls or urlparse(url).scheme not in {"http", "https"}:
            continue
        seen_urls.add(url)
        title = _clean(source.get("title"), 300) or url
        if _contains_price_value(title):
            title = urlparse(url).netloc or "Official product source"
        sections["warranty"].append({"label": f"Source {len(seen_urls)}", "value": title, "url": url})
    if not any(sections.values()):
        raise ValueError("No supported specification facts were returned.")
    return sections


def structure_specification(computer_spec: str, answer: str, sources: list[dict]) -> dict[str, list[dict[str, str]]]:
    from services.ollama_client import OllamaClient
    from services.prompt_service import render_prompt
    from services.settings_service import get_setting, get_task_model

    source_text = "\n".join(
        f"[{index}] {source.get('title') or source.get('url')}: {source.get('url')}"
        for index, source in enumerate(sources, start=1)
    )
    prompt = render_prompt(
        "spec_sheet_structuring",
        computer_spec=computer_spec,
        research_answer=answer,
        sources=source_text or "No source links supplied.",
    )
    client = OllamaClient(get_setting("ollama_url") or "")
    model = get_setting("computer_finder_model") or get_task_model("chat_answering") or "llama3.2"
    parsed, raw_response, error = client.generate_json(model, prompt)
    if parsed is None or error:
        raise ValueError(f"Could not structure the specification result: {error or raw_response}")
    return _normalise_sections(parsed, sources)


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    colour = OxmlElement("w:color")
    colour.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([colour, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _set_cell(cell, text: str, url: str = "") -> None:
    paragraph = cell.paragraphs[0]
    _clear_paragraph(paragraph)
    if url:
        _add_hyperlink(paragraph, text, url)
    else:
        paragraph.add_run(text)


def _populate_table(table, rows: list[dict[str, str]]) -> None:
    template_row = deepcopy(table.rows[1]._tr)
    for row in list(table.rows)[1:]:
        table._tbl.remove(row._tr)
    for row_data in rows:
        row_xml = deepcopy(template_row)
        table._tbl.append(row_xml)
        row = table.rows[-1]
        _set_cell(row.cells[0], row_data["label"])
        _set_cell(row.cells[1], row_data["value"], row_data.get("url", ""))
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)


def build_specification_document(sections: dict[str, list[dict[str, str]]]) -> bytes:
    if not SPECIFICATION_TEMPLATE.exists():
        raise FileNotFoundError("The machine specification template is missing.")
    document = Document(str(SPECIFICATION_TEMPLATE))
    for table, key in zip(document.tables, SECTION_KEYS, strict=True):
        _populate_table(table, sections.get(key) or [])
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def specification_filename(sections: dict[str, list[dict[str, str]]]) -> str:
    model = next(
        (row["value"] for row in sections.get("system", []) if row["label"].casefold() == "manufacturer and model"),
        "Machine Specification",
    )
    safe_name = re.sub(r"[^A-Za-z0-9]+", "_", model).strip("_")[:100] or "Machine_Specification"
    return f"Machine_Specification_{safe_name}.docx"


def specification_text(sections: dict[str, list[dict[str, str]]]) -> str:
    headings = {
        "system": "Quoted system",
        "connectivity": "Connectivity and peripherals",
        "warranty": "Warranty and product information",
    }
    parts = ["# Machine Specification"]
    for key in SECTION_KEYS:
        parts.extend(["", f"## {headings[key]}", ""])
        parts.extend(f"- **{row['label']}:** {row['value']}" for row in sections.get(key, []))
    return "\n".join(parts).strip() + "\n"


def generate_specification_document(computer_spec: str, answer: str, sources: list[dict]) -> tuple[bytes, str, str]:
    sections = structure_specification(computer_spec, answer, sources)
    return build_specification_document(sections), specification_filename(sections), specification_text(sections)
