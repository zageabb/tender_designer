from __future__ import annotations

import csv
from email import policy
from email.parser import BytesParser
from pathlib import Path


def _markdown_escape(value) -> str:
    text = str(value or "").replace("\n", " ").strip()
    return text.replace("|", "\\|")


def _rows_to_markdown(rows: list[list[str]]) -> str:
    cleaned_rows = [[_markdown_escape(cell) for cell in row] for row in rows if any(str(cell).strip() for cell in row)]
    if not cleaned_rows:
        return ""
    column_count = max(len(row) for row in cleaned_rows)
    padded_rows = [row + [""] * (column_count - len(row)) for row in cleaned_rows]
    header = padded_rows[0]
    separator = ["---"] * column_count
    body = padded_rows[1:]
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(separator) + " |",
        *["| " + " | ".join(row) + " |" for row in body],
    ]
    return "\n".join(lines)


def _docx_to_markdown(path: Path) -> str:
    from docx import Document

    document = Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "").lower()
        if style_name.startswith("heading"):
            level_text = "".join(character for character in style_name if character.isdigit())
            level = min(max(int(level_text or "1"), 1), 6)
            lines.append(f"{'#' * level} {text}")
        elif "list bullet" in style_name:
            lines.append(f"- {text}")
        elif "list number" in style_name:
            lines.append(f"1. {text}")
        else:
            lines.append(text)

    for table in document.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        table_markdown = _rows_to_markdown(rows)
        if table_markdown:
            if lines:
                lines.append("")
            lines.append(table_markdown)

    return "\n".join(lines)


def _csv_to_markdown(path: Path) -> str:
    with path.open("r", encoding="utf-8", errors="ignore", newline="") as handle:
        rows = list(csv.reader(handle))
    return _rows_to_markdown(rows)


def _xlsx_to_markdown(path: Path) -> str:
    from openpyxl import load_workbook

    workbook = load_workbook(filename=str(path), data_only=True)
    sections: list[str] = []
    for sheet in workbook.worksheets:
        rows: list[list[str]] = []
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() if value is not None else "" for value in row]
            if any(values):
                rows.append(values)
        sheet_table = _rows_to_markdown(rows)
        if sheet_table:
            sections.append(f"## Sheet: {sheet.title}\n\n{sheet_table}")
        else:
            sections.append(f"## Sheet: {sheet.title}\n\n_No populated rows found._")
    return "\n\n".join(sections)


def _eml_to_markdown(path: Path) -> str:
    with path.open("rb") as handle:
        message = BytesParser(policy=policy.default).parse(handle)
    parts = [
        "# Email",
        "",
        f"- Subject: {message.get('subject', '').strip()}",
        f"- From: {message.get('from', '').strip()}",
        f"- To: {message.get('to', '').strip()}",
    ]
    body = message.get_body(preferencelist=("plain", "html"))
    if body:
        parts.extend(["", "## Body", "", body.get_content().strip()])
    return "\n".join(parts).strip()


def extract_text(file_path: str | Path) -> tuple[str, str | None]:
    path = Path(file_path)
    suffix = path.suffix.lower()
    try:
        if suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore"), None
        if suffix == ".csv":
            return _csv_to_markdown(path), None
        if suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
            return text, None
        if suffix == ".docx":
            return _docx_to_markdown(path), None
        if suffix == ".xlsx":
            return _xlsx_to_markdown(path), None
        if suffix == ".eml":
            return _eml_to_markdown(path), None
        if suffix == ".msg":
            return "", "MSG extraction is not yet configured in this initial version."
        return "", f"Unsupported file type: {suffix}"
    except Exception as exc:
        return "", str(exc)
