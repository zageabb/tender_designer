from io import BytesIO

from docx import Document

from services.specification_document import (
    _normalise_sections,
    build_specification_document,
    specification_filename,
)


def test_specification_document_uses_dynamic_rows_and_excludes_prices():
    sections = _normalise_sections(
        {
            "sections": {
                "system": [
                    {"label": "Manufacturer and model", "value": "Example ProBook X1"},
                    {"label": "Processor", "value": "Example CPU, 12 cores"},
                    {"label": "Price", "value": "£999"},
                    {"label": "Availability", "value": "GBP 999 from supplier"},
                ],
                "connectivity": [{"label": "Wired networking", "value": "2.5GbE"}],
                "warranty": [{"label": "Warranty", "value": "3-year onsite warranty"}],
            }
        },
        [{"title": "Example ProBook X1 - $999", "url": "https://example.com/product"}],
    )

    content = build_specification_document(sections)
    document = Document(BytesIO(content))
    text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)

    assert "Example ProBook X1" in text
    assert "2.5GbE" in text
    assert "3-year onsite warranty" in text
    assert "Lenovo ThinkCentre" not in text
    assert "£999" not in text
    assert "GBP 999" not in text
    assert "$999" not in text
    assert "example.com" in text
    assert specification_filename(sections) == "Machine_Specification_Example_ProBook_X1.docx"


def test_specification_document_source_link_is_clickable():
    sections = _normalise_sections(
        {
            "sections": {
                "system": [{"label": "Manufacturer and model", "value": "Example Workstation"}],
                "connectivity": [],
                "warranty": [],
            }
        },
        [{"title": "Official product page", "url": "https://example.com/specification"}],
    )

    document = Document(BytesIO(build_specification_document(sections)))
    relationships = document.part.rels.values()

    assert any(rel.target_ref == "https://example.com/specification" for rel in relationships)
